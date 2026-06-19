"""Generate interview-level Word document for Fragmented Consciousness project."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\coding folders\fragmented-consciousness-complete\fragmented-consciousness\Fragmented-Consciousness-Interview-Guide.docx"


def set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x5C)


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)


def add_h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x6B, 0x8A)
        run.font.size = Pt(13)


def add_h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_defaults(doc)

    # Margins for ~3 pages
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    add_title(doc, "Fragmented Consciousness")
    add_subtitle(doc, "Distributed Artificial Lifeform — Complete Technical & Interview Guide")
    doc.add_paragraph()

    # --- PAGE 1: OVERVIEW ---
    add_h1(doc, "1. Project Overview")
    doc.add_paragraph(
        "Fragmented Consciousness is a browser-based multi-window experiment where a single 3D humanoid "
        "exists across six synchronized popup windows. Each window represents a cognitive subsystem "
        "(Body, Vision, Memory, Touch, Hearing, Emotion). Consciousness level (1–6) depends on how many "
        "panes are open. The creature chases the user's cursor, learns jumps and walls, pushes its window "
        "toward off-screen cursors, and can briefly grab the cursor when fully conscious (6/6)."
    )

    add_h2(doc, "Core Concept")
    add_bullets(doc, [
        "Consciousness is distributed — no single window holds the full mind.",
        "Body = somatic core (Three.js 3D simulation); other panes = lightweight 2D sensor canvases.",
        "BroadcastChannel API syncs state across windows without a backend server.",
        "Launcher page also broadcasts desktop pointer so cursor on dashboard steers the body.",
    ])

    add_h2(doc, "Technology Stack")
    add_table(doc, ["Layer", "Technology", "Purpose"], [
        ["Frontend", "React 19 + TypeScript", "UI, window pages, hooks"],
        ["3D", "Three.js r184", "Humanoid rendering, grid floor, particles"],
        ["Styling", "Tailwind CSS 4", "Sci-fi glassmorphism UI"],
        ["Routing", "Wouter", "SPA routes (/window/body, etc.)"],
        ["Build", "Vite 7", "Dev server port 3000, production bundle"],
        ["Server", "Express + esbuild", "Serves dist/public in production"],
        ["Sync", "BroadcastChannel", "Cross-window creature, memory, sensorium"],
        ["Hand track", "MediaPipe Hand Landmarker", "Optional webcam fingertip control"],
        ["Testing", "Vitest", "Unit tests for coords, physics, gating"],
        ["Model", "CesiumMan glTF (Khronos)", "Humanoid mesh + run animation"],
    ])

    add_h2(doc, "Architecture (Key Files)")
    add_bullets(doc, [
        "sharedState.ts — SharedStateManager, presence heartbeats, consciousness recompute",
        "creatureBrain.ts — Pursuit physics, jump, grab, window-push state machine",
        "subsystemGates.ts — Capability matrix from open panes (speed, sight, learning)",
        "desktopCoords.ts — Screen ↔ Body canvas coordinate mapping",
        "playSpace.ts — Normalized 0..1 play bounds, render space conversion",
        "threeSetup.ts — Orthographic top-down camera, creature, grid, cursor hint",
        "WindowBody.tsx — Main simulation loop (~60 FPS)",
        "useDesktopPointerBroadcast.ts — Mouse screenX/Y → sensorium.desktopPointer",
    ])

    doc.add_page_break()

    # --- PAGE 2: ISSUES & FIXES + PHYSICS ---
    add_h1(doc, "2. Issues Faced & How We Rectified Them")
    add_table(doc, ["Issue", "Root Cause", "Fix"], [
        ["Character invisible / off-screen", "Wrong coordinate space; figure at y≈912; broken multiWindow ortho scene", "Reverted to per-pane Body + individual 2D panes; normToRender centered on canvas"],
        ["Cursor glue instead of chase", "Direct position snap to mouse", "Velocity-based pursuit in creatureBrain.ts (accel, friction, maxSpeed)"],
        ["Jitter at center / instant grab", "Vision coords used as Body coords; grab on hover", "desktopCoords mapping; sustained proximity grab; cursor.active required"],
        ["No response outside Body tab", "Each popup only gets mousemove over itself; Launcher not broadcasting", "desktopPointer on all panes + Launcher; last-known pointer 8s TTL"],
        ["3/6 confused behavior", "Partial targetWeight = (level/6)²; missing Touch/Hearing/Emotion", "By design — document Open All Panes; impairment HUD hints"],
        ["Character walks off visible area", "Perspective camera slanted; no follow cam", "Orthographic top-down camera + updateCameraFocus follows creature"],
        ["Invisible wall jitter", "Hard bounce at bounds.min/max; memory repel teleport", "Slide-along-edge velocity; repel as force not teleport"],
        ["Push but no capture after", "window_push only called moveTo, no lunge", "During push: if cursor enters pane → sprint + fast grab (0.28s)"],
        ["Touch click no visible flinch", "Pulse coords not mapped to Body space", "mapPaneNormToContainer for touch/hearing pulses"],
        ["Windows deploy failed", "NODE_ENV=production Unix-only in package.json", "node dist/index.js + fs.existsSync for static path"],
    ])

    add_h2(doc, "Physics & Simulation (creatureBrain.ts)")
    doc.add_paragraph(
        "Simulation runs in normalized tab space: (0,0) = top-left of Body canvas, (1,1) = bottom-right. "
        "Three.js render coords derived via playSpace.normToRender(). Fixed timestep capped at 50ms per frame."
    )
    add_bullets(doc, [
        "Chase: acceleration toward clamped target edge when cursor is off-screen (edgeBoost 1.45×)",
        "Velocity integration: vel += accel×dt; clamp to maxSpeed; apply friction (0.88–0.93)",
        "Gravity: GRAVITY=22 for jumps; landing updates jumpAffinity learning",
        "Bounds: soft wall — zero outward velocity, slide along edge (not bounce)",
        "Modes: chase | jumping | grabbed | window_push | release",
        "Grab: CATCH_RADIUS=0.028, GRAB_SUSTAIN=0.65s, GRAB_LUNGE=0.28s during window push",
        "Window push: window.moveTo() nudges popup toward cursor.screenX/Y (needs Hearing+Memory+4/6)",
    ])

    add_h2(doc, "How the Character Learns")
    add_bullets(doc, [
        "Jump learning (Memory pane): jumpAffinity 0–1; successful jumps (preDist − landedDist > 0.04) increase affinity; failed jumps decrease it",
        "Wall awareness (Memory + Touch): wallStreak on boundary hits; wallAwareness 0–1 unlocks window push",
        "Memory paths: Body writes safe/danger/exploration paths to sharedState; Memory pane visualizes; danger points repel creature via memoryRepel force",
        "Window push learned: windowPushLearned flag after push mode completes",
        "Grab count: skills.grabs incremented on successful capture (Touch + 6/6 required)",
        "No ML/neural net — learning is heuristic state accumulation (interview: ' lightweight reinforcement-style skills')",
    ])

    doc.add_page_break()

    # --- PAGE 3: SUBSYSTEMS, COORDS, INTERVIEW ---
    add_h1(doc, "3. Subsystems, Coordinates & Interview Talking Points")
    add_table(doc, ["Pane", "Unlocks", "Broadcasts"], [
        ["Body", "3D sim, physics loop, creature state", "desktopPointer, creature updates, memory paths"],
        ["Vision", "Sight (targetWeight > 0)", "visionTarget (pane 0–1), desktopPointer"],
        ["Memory", "Jump learn, danger repel, wall learn", "desktopPointer"],
        ["Touch", "Flinch on click, grab at 6/6", "touchPulse, desktopPointer"],
        ["Hearing", "Sound pull, window nudge eligibility", "soundPulse, desktopPointer"],
        ["Emotion", "Speed/jump multiplier, mood tint", "desktopPointer"],
    ])

    add_h2(doc, "Coordinate Pipeline")
    add_numbered(doc, [
        "User moves mouse → pane broadcasts desktopPointer {x: screenX, y: screenY}",
        "Body reads pointer → screenToContainerNorm() maps to Body canvas (rawNx/rawNy can be <0 or >1)",
        "resolvePerceivedTarget() blends wander + vision with smoothed lerp",
        "stepCreature() chases toward target; clamped edge point when off-screen",
        "normToRender() → Three.js position; updateCameraFocus() keeps character in frame",
    ])

    add_h2(doc, "Models & External Assets")
    add_bullets(doc, [
        "CesiumMan glTF — client/public/models/cesium-man.glb (Khronos sample humanoid, ~1.8 unit height scaled)",
        "Procedural placeholder capsule humanoid while glTF loads",
        "MediaPipe hand_landmarker.task — loaded from Google CDN for optional hand tracking",
        "GLTFLoader + AnimationMixer for run cycle playback",
    ])

    add_h2(doc, "Consciousness Formula (subsystemGates.ts)")
    add_bullets(doc, [
        "level = count of open subsystem types (including Body)",
        "complete = level / 6",
        "targetWeight = hasVision ? complete² : 0  (blind without Vision)",
        "maxSpeed = 0.05 + complete×0.47 | accel = 0.15 + complete×1.5",
        "canGrab = Touch open AND complete >= 1 (all six panes)",
        "canWindowPush = Hearing + Memory + complete >= 0.67",
    ])

    add_h2(doc, "Interview Q&A — Advanced")
    add_bullets(doc, [
        "Q: Why BroadcastChannel over WebSocket? A: Same-origin popups, zero server, instant sync, no infra cost.",
        "Q: Why normalized 0..1 space? A: Resolution-independent; tab resize recomputes playBounds; same sim on any monitor.",
        "Q: How handle cursor outside window? A: Map screen coords to Body space; chase to edge; push window; lunge on entry.",
        "Q: Coupling vs capabilities? A: Overlap/coupling is visual only; pursuit gates use pane presence count.",
        "Q: Test strategy? A: Vitest unit tests for desktopCoords, playSpace bounds, stepCreature edge chase, grab guards.",
    ])

    add_h2(doc, "Run, Build & Deploy")
    add_bullets(doc, [
        "Dev: pnpm dev → http://localhost:3000 → Open All Panes",
        "Test: pnpm test | Typecheck: pnpm check",
        "Build: pnpm build → dist/public + dist/index.js",
        "Prod: pnpm start → Express serves SPA on port 3000",
    ])

    add_h2(doc, "Legacy / Unused Code (mention if asked)")
    add_bullets(doc, [
        "WindowManager.ts, worldState.ts — localStorage desktop sync (superseded by BroadcastChannel)",
        "multiWindowScene.ts — abandoned orthographic multi-window Three.js experiment",
    ])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
